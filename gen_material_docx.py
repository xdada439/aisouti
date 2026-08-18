from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('编程与计算机基础 - 复习资料', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Chapter 1
doc.add_heading('第一章：编程基础', level=1)

doc.add_heading('1.1 Java基础', level=2)
p = doc.add_paragraph()
p.add_run('Java是一种面向对象的编程语言。Java的基本数据类型包括：').font.size = Pt(12)
p = doc.add_paragraph('• int：整数类型，用于存储整数值，如 1, 100, -50', style='List Bullet')
p = doc.add_paragraph('• double：双精度浮点数类型', style='List Bullet')
p = doc.add_paragraph('• boolean：布尔类型，值为 true 或 false', style='List Bullet')
p = doc.add_paragraph('• char：字符类型，存储单个字符', style='List Bullet')
p = doc.add_paragraph('• byte、short、long、float 等', style='List Bullet')
doc.add_paragraph('注意：String不是基本数据类型，它是Java中的一个类，属于引用类型。ArrayList和HashMap也属于Java集合框架中的类，不是基本数据类型。')

doc.add_paragraph('Java中，一个类可以实现多个接口，这是Java实现多继承的方式。所有类的父类是Object类，它定义了equals()、hashCode()、toString()等基础方法。')

doc.add_heading('1.2 面向对象编程', level=2)
doc.add_paragraph('面向对象编程的三大特性是：封装、继承和多态。')
doc.add_paragraph('• 封装：将数据和操作数据的方法绑定在一起，隐藏内部实现细节。通过private修饰符和public的getter/setter方法实现。', style='List Bullet')
doc.add_paragraph('• 继承：子类可以继承父类的属性和方法，实现代码复用。Java使用extends关键字实现继承。', style='List Bullet')
doc.add_paragraph('• 多态：同一个方法调用可以有不同的行为。通过方法重载（编译时多态）和方法重写（运行时多态）实现。', style='List Bullet')
doc.add_paragraph('递归是函数调用自身的技术，不属于面向对象三大特性。')

doc.add_heading('1.3 Python基础', level=2)
doc.add_paragraph('Python是一种解释型、动态类型的编程语言。定义函数使用def关键字：')
doc.add_paragraph('def my_function():', style='List Bullet')
doc.add_paragraph('    print("Hello World")', style='List Bullet')
doc.add_paragraph('Python中列表(list)是可变的，可以修改其中的元素；元组(tuple)是不可变的，创建后不能修改。')

doc.add_heading('1.4 Kotlin基础', level=2)
doc.add_paragraph('Kotlin是现代Android开发的首选语言。声明不可变变量可以使用val或const val关键字：')
doc.add_paragraph('• val：声明只读变量，运行时赋值后不可修改', style='List Bullet')
doc.add_paragraph('• const val：编译时常量，必须在编译时确定值', style='List Bullet')
doc.add_paragraph('• var：声明可变变量，可以重新赋值', style='List Bullet')
doc.add_paragraph('• let：是Kotlin的作用域函数，不是变量声明关键字', style='List Bullet')

# Chapter 2
doc.add_heading('第二章：网络通信', level=1)

doc.add_heading('2.1 HTTP协议', level=2)
doc.add_paragraph('HTTP（超文本传输协议）是互联网应用最广泛的协议之一。HTTP协议默认使用80端口进行通信，HTTPS使用443端口。')
doc.add_paragraph('常见的HTTP请求方法包括：')
doc.add_paragraph('• GET：获取资源，请求参数在URL中', style='List Bullet')
doc.add_paragraph('• POST：提交数据，数据在请求体中，适合传输大量或敏感数据', style='List Bullet')
doc.add_paragraph('• PUT：更新资源（完整替换）', style='List Bullet')
doc.add_paragraph('• DELETE：删除资源', style='List Bullet')
doc.add_paragraph('• PATCH：部分更新资源', style='List Bullet')
doc.add_paragraph('• HEAD、OPTIONS、TRACE等', style='List Bullet')
doc.add_paragraph('注意：SEND不是标准的HTTP请求方法。')

doc.add_heading('2.2 TCP协议', level=2)
doc.add_paragraph('TCP（传输控制协议）是一种面向连接的、可靠的传输层协议。与UDP不同，TCP需要建立连接（三次握手）才能传输数据，保证数据的有序性和完整性。因此，声称TCP是面向无连接协议的说法是错误的。')

# Chapter 3
doc.add_heading('第三章：数据库', level=1)
doc.add_heading('3.1 关系型数据库', level=2)
doc.add_paragraph('关系型数据库以表格形式存储数据，使用SQL语言进行操作。常见的关系型数据库包括：')
doc.add_paragraph('• MySQL：开源的关系型数据库，广泛用于Web应用', style='List Bullet')
doc.add_paragraph('• PostgreSQL：功能强大的开源数据库，支持复杂查询和地理空间数据', style='List Bullet')
doc.add_paragraph('• Oracle：商业数据库，用于大型企业系统', style='List Bullet')
doc.add_paragraph('• SQL Server：微软的关系型数据库', style='List Bullet')
doc.add_paragraph('注意：MongoDB是非关系型（NoSQL）文档数据库，Redis是非关系型键值存储数据库。')

# Chapter 4
doc.add_heading('第四章：Android开发', level=1)

doc.add_heading('4.1 Activity生命周期', level=2)
doc.add_paragraph('Activity是Android应用的基本组件之一。其生命周期回调方法调用顺序为：')
doc.add_paragraph('onCreate() → onStart() → onResume() → (Activity运行中) → onPause() → onStop() → onDestroy()')
doc.add_paragraph('因此onCreate之后紧接着调用的是onStart()，而不是onResume()或其他方法。')

doc.add_heading('4.2 Android四大组件', level=2)
doc.add_paragraph('Android的四大组件包括：')
doc.add_paragraph('• Activity：用户界面组件，每个屏幕对应一个Activity', style='List Bullet')
doc.add_paragraph('• Service：后台服务组件，用于执行长时间运行的操作。注意：Service默认运行在主线程中，而不是子线程。如需在子线程执行任务，需要使用IntentService或手动创建线程。', style='List Bullet')
doc.add_paragraph('• BroadcastReceiver：广播接收器，用于接收系统和应用广播', style='List Bullet')
doc.add_paragraph('• ContentProvider：内容提供者，用于应用间数据共享', style='List Bullet')
doc.add_paragraph('Fragment不是四大组件之一，它是Activity中的UI片段。')

doc.add_heading('4.3 数据存储', level=2)
doc.add_paragraph('Android提供了多种数据存储方式。SharedPreferences是用于存储简单键值对数据的轻量级存储类，适合保存用户设置、登录状态等少量数据。它基于XML文件存储，数据以键值对的形式组织。')

# Chapter 5
doc.add_heading('第五章：版本控制', level=1)
doc.add_heading('5.1 Git基础', level=2)
doc.add_paragraph('Git是目前最流行的分布式版本控制系统。常用命令包括：')
doc.add_paragraph('• git log：查看提交历史记录，显示提交的哈希值、作者、日期和提交信息', style='List Bullet')
doc.add_paragraph('• git status：查看工作区和暂存区的状态', style='List Bullet')
doc.add_paragraph('• git diff：查看文件修改的详细差异', style='List Bullet')
doc.add_paragraph('• git branch：查看和管理分支', style='List Bullet')
doc.add_paragraph('git pull命令的作用是从远程仓库获取最新代码并合并到本地分支。它等同于先执行git fetch（获取远程更新），再执行git merge（合并到当前分支）。')

# Chapter 6
doc.add_heading('第六章：数据结构与算法', level=1)
doc.add_heading('6.1 排序算法', level=2)
doc.add_paragraph('排序算法是计算机科学中的基础算法。不同算法的时间复杂度不同：')
doc.add_paragraph('• 快速排序：平均时间复杂度O(nlogn)，最坏O(n²)，是最常用的高效排序算法', style='List Bullet')
doc.add_paragraph('• 归并排序：时间复杂度稳定为O(nlogn)，需要额外空间', style='List Bullet')
doc.add_paragraph('• 冒泡排序：时间复杂度O(n²)，仅适合教学使用', style='List Bullet')
doc.add_paragraph('• 插入排序：时间复杂度O(n²)，对小规模或基本有序的数据较好', style='List Bullet')
doc.add_paragraph('• 选择排序：时间复杂度O(n²)，无论数据如何都执行固定次数的比较', style='List Bullet')

# Chapter 7
doc.add_heading('第七章：Linux基础', level=1)
doc.add_heading('7.1 常用命令', level=2)
doc.add_paragraph('Linux系统中常用的命令包括：')
doc.add_paragraph('• chmod：修改文件或目录的权限（change mode）', style='List Bullet')
doc.add_paragraph('• chown：修改文件或目录的所有者（change owner）', style='List Bullet')
doc.add_paragraph('• chgrp：修改文件或目录的所属组', style='List Bullet')
doc.add_paragraph('• ls：列出目录内容', style='List Bullet')
doc.add_paragraph('• pwd：显示当前工作目录的完整路径（print working directory）', style='List Bullet')
doc.add_paragraph('• cd：切换目录', style='List Bullet')

# Save
path = r'D:\AIProjects\ai-souti\test_material.docx'
doc.save(path)
print(f'Saved: {path}')
